"""
DOCX Generator — crea un archivo Word formateado con el guión generado.
"""
import io
import re
from datetime import datetime


def generate_docx(script_text: str, titulo: str = "Guión", canal: str = "") -> bytes:
    """
    Genera un archivo .docx con el guión formateado.
    Retorna bytes del archivo.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Estilos globales ──────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Portada ───────────────────────────────────────────────────
    titulo_para = doc.add_paragraph()
    titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_para.add_run(titulo)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    fecha_para = doc.add_paragraph()
    fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_run = fecha_para.add_run(datetime.now().strftime("%d de %B de %Y"))
    fecha_run.font.size = Pt(11)
    fecha_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if canal:
        canal_para = doc.add_paragraph()
        canal_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        canal_run = canal_para.add_run(f"Canal: {canal}")
        canal_run.font.size = Pt(10)
        canal_run.italic = True
        canal_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # espaciado
    doc.add_page_break()

    # ── Contenido del guión ───────────────────────────────────────
    # Parsear secciones: **[NOMBRE]** como heading
    SECTION_RE = re.compile(r"\*\*\[([^\]]+)\]\*\*")
    TITLE_RE   = re.compile(r"\*\*TÍTULOS SUGERIDOS.*?\*\*", re.IGNORECASE)

    lines = script_text.split("\n")
    current_section = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Línea separadora
        if stripped == "---":
            doc.add_paragraph()
            continue

        # Encabezado de sección
        sec_match = SECTION_RE.match(stripped)
        if sec_match:
            current_section = sec_match.group(1)
            h = doc.add_heading(current_section, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            continue

        # Bloque de títulos sugeridos
        if TITLE_RE.search(stripped):
            h = doc.add_heading("Títulos Sugeridos", level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            continue

        # Numeración de títulos (1. / 2. / 3.)
        if re.match(r"^\d+\.", stripped) and current_section is None:
            p = doc.add_paragraph(stripped, style="List Number")
            continue

        # Texto normal
        p = doc.add_paragraph()
        # Negrita inline **texto**
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    # ── Footer (canal + fecha) ─────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = f"{canal} — Generado el {datetime.now().strftime('%d/%m/%Y')}" if canal else f"Generado el {datetime.now().strftime('%d/%m/%Y')}"
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer_run.italic = True

    # ── Serializar a bytes ─────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
